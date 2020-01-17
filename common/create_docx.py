from docx import Document
from docx.shared import Pt,Inches,Cm
import pandas as pd
import datetime
from datetime import timedelta
import math
from pyecharts.charts import Line,Bar, Page,Kline,Grid
from pyecharts import options as opts
from pyecharts.render import make_snapshot
from snapshot_selenium import snapshot
from common.plotbypyecharts import kline_profession


from common.db_operation import mysql_login

# 行情明细脚本
sql = '''
SELECT  DateTime
        ,PRE_CLOSE
        ,OPEN
        ,CLOSE
        ,HIGH
        ,LOW
        ,VOLUME
        ,AMT
        ,PCT_CHG
        ,SWING
FROM    ext_data_stock.stock_quotation_his
'''

# 昨日行情模板

temp_yest = '''
前一交易日（{last_trade_date}）上证综指开盘{open_price}，最高{high_price}，\
最低{low_price}，振幅{swing}%，收盘{close_price}，相对前收盘点位（{pre_close_price}）涨跌幅为{pct_chg}%，\
其中成交量为{trade_vol}万股，成交金额{trade_amt}亿元。\
'''

# 区间行情模板

temp_period = '''
近{days}天内（包含{trade_days}个交易日），上证综指最高达到{high_price_period}，最低降至{low_price_period}，\
平均点位{mean_price_period}，期间涨跌幅{pct_chg_period}%。\
'''


# 获取单日行情数据
def get_quot_day():
    #创建连接
    conn = mysql_login()
    # pandas查询
    df = pd.read_sql(sql=sql, con=conn)
    conn.close()
    # 最新行情
    data_curr = df.iloc[-1]
    last_trade_date = data_curr['DateTime'].strftime("%Y-%m-%d")
    open_price = round(data_curr['OPEN'],0)
    close_price = round(data_curr['CLOSE'],0)
    high_price = round(data_curr['HIGH'],0)
    low_price = round(data_curr['LOW'],0)
    swing = round(data_curr['SWING'],2)
    pre_close_price = round(data_curr['PRE_CLOSE'],0)
    pct_chg = round(data_curr['PCT_CHG'],2)
    trade_vol = round(data_curr['VOLUME']/1e4,2)
    trade_amt = round(data_curr['AMT']/1e8,2)

    ph_curr = temp_yest.format_map(vars())
    return ph_curr


# 获取区间行情数据
def get_quot_period(days):
    #创建连接
    conn = mysql_login()
    # pandas查询
    df = pd.read_sql(sql=sql, con=conn)
    conn.close()
    # 最新行情
    data_curr = df.iloc[-1]
    # 近90天行情
    start_date = data_curr['DateTime'] - timedelta(days=days)
    end_date = data_curr['DateTime']
    data_period = df[(df['DateTime']>=start_date) & (df['DateTime']<=end_date)]

    picture = plot_k_picture(data_period)
    # picture.render('../result_set/000001SH.png')

    days = days
    trade_days = int(data_period['CLOSE'].describe()['count'])
    low_price_period = round(data_period['CLOSE'].describe()['min'],0)
    high_price_period = round(data_period['CLOSE'].describe()['max'],0)
    mean_price_period = round(data_period['CLOSE'].describe()['mean'],0)
    pct_chg_period = round(calcu_stock_range_period(data_period,start_date,end_date)*100,2)
    
    ph_period = temp_period.format_map(vars())
    return ph_period, data_period, picture


# 计算股票区间涨跌幅
def calcu_stock_range_period(df, start_date, end_date):
    data = df[(df['DateTime'] >= start_date) & (df['DateTime'] <= end_date) & (df['VOLUME'] > 0)]
    stock_range_period = math.exp((data['CLOSE']/data['PRE_CLOSE']).apply(lambda x: math.log(x)).sum()) - 1
    return stock_range_period


# 画日k线图
def plot_k_picture(data_period):
    return kline_profession(data_period)


def output_docx(doc_path='../result_set/doc_test1.docx', *p_days):
    # 创建文档
    document = Document()

    # 获取最新行情统计
    ph_curr = get_quot_day()

    # 新建段落,写入数据
    paragraph = document.add_paragraph(ph_curr)
    ph_format = paragraph.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)   # 设置段后间距
    ph_format.line_spacing = Pt(19)  # 设置行间距

    # 获取最近200天行情统计信息
    ph_period, data_dtl, picture = get_quot_period(200)
    picture.render('../result_set/000001SH.html')
    make_snapshot(snapshot, picture.render(), '../result_set/000001SH.png')

    # 新建段落,写入数据
    paragraph = document.add_paragraph(ph_period)
    ph_format = paragraph.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.line_spacing = Pt(19)  # 设置行间距

    # 新建表格1
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    # 插入图片
    run = paragraph.add_run()
    run.add_picture('../result_set/000001SH.png', width=Inches(6.0))

    # 新建表格2
    data_dtl = data_dtl[['DateTime', 'OPEN', 'CLOSE', 'LOW', 'HIGH', 'VOLUME', 'AMT', 'PCT_CHG']]
    data_dtl['DateTime'] = data_dtl['DateTime'].apply(lambda x: x.strftime("%Y-%m-%d"))

    table2 = document.add_table(rows=len(data_dtl.index), cols=len(data_dtl.columns))
    table2.add_row()
    for i in range(len(data_dtl.columns)):
        table2.cell(0, i).text = data_dtl.columns[i]  # 添加表头

    for row in range(1, len(data_dtl.index)+1):
        for col in range(len(data_dtl.columns)):
            table2.cell(row, col).width = 1
            table2.cell(row, col).text = str(data_dtl.iloc[row-1, col])
            table2.cell(row, col).width = Cm(6)
    table2.style = 'Medium Grid 1 Accent 1'
    table2.autofit = True

    # 保存文档
    document.save(doc_path)


if __name__ == "__main__":
    output_docx('../result_set/doc_test1.docx')

